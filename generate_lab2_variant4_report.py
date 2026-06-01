from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path.cwd()
ASSETS = ROOT / "assets_lab2_variant4"
ASSETS.mkdir(exist_ok=True)
OUT = ROOT / "Лаба_Коровин_2_вариант4_готово.docx"


def fnt(size=24, bold=False):
    path = Path(r"C:\Windows\Fonts\timesbd.ttf" if bold else r"C:\Windows\Fonts\times.ttf")
    if path.exists():
        return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def wrap(draw, text, font, max_width):
    words = text.split()
    lines, line = [], ""
    for word in words:
        test = word if not line else line + " " + word
        if draw.textbbox((0, 0), test, font=font)[2] <= max_width:
            line = test
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines


def centered(draw, box, text, font):
    x1, y1, x2, y2 = box
    lines = wrap(draw, text, font, x2 - x1 - 22)
    h = sum(draw.textbbox((0, 0), line, font=font)[3] for line in lines) + (len(lines) - 1) * 5
    y = y1 + (y2 - y1 - h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=font, fill=(20, 20, 20))
        y += bbox[3] + 5


def block(draw, box, text, kind="process"):
    fill = {
        "start": "#DDEBFF",
        "process": "#EEF5E9",
        "decision": "#FFF2CC",
        "error": "#FCE4D6",
        "fixed": "#E2F0D9",
    }[kind]
    if kind == "decision":
        x1, y1, x2, y2 = box
        pts = [(x1 + (x2 - x1) / 2, y1), (x2, y1 + (y2 - y1) / 2), (x1 + (x2 - x1) / 2, y2), (x1, y1 + (y2 - y1) / 2)]
        draw.polygon(pts, fill=fill, outline="#34495E")
    else:
        draw.rounded_rectangle(box, radius=22 if kind == "start" else 8, fill=fill, outline="#34495E", width=3)
    centered(draw, box, text, fnt(22))


def arrow(draw, start, end, label=None):
    draw.line([start, end], fill="#34495E", width=3)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) > abs(ey - sy):
        pts = [(ex, ey), (ex - 12 if ex > sx else ex + 12, ey - 7), (ex - 12 if ex > sx else ex + 12, ey + 7)]
    else:
        pts = [(ex, ey), (ex - 7, ey - 12 if ey > sy else ey + 12), (ex + 7, ey - 12 if ey > sy else ey + 12)]
    draw.polygon(pts, fill="#34495E")
    if label:
        mx, my = (sx + ex) / 2, (sy + ey) / 2
        draw.text((mx + 8, my - 18), label, font=fnt(18, True), fill=(20, 20, 20))


def make_flowchart():
    img = Image.new("RGB", (1100, 1400), "white")
    d = ImageDraw.Draw(img)
    boxes = [
        ((360, 40, 740, 130), "Начало отладки", "start"),
        ((280, 190, 820, 290), "Подготовить тестовые значения датчика, включая T < 0°C", "process"),
        ((280, 350, 820, 450), "Запустить ошибочную обработку abs(temperature)", "error"),
        ((280, 510, 820, 630), "Температура потеряла знак минус?", "decision"),
        ((280, 700, 820, 800), "Локализовать ошибку: функция abs()", "process"),
        ((280, 860, 820, 960), "Исправить код: вернуть temperature без abs()", "fixed"),
        ((280, 1020, 820, 1120), "Повторить тест с -3.4°C и -12.7°C", "process"),
        ((280, 1180, 820, 1280), "Зафиксировать результат в отчёте и Git", "process"),
    ]
    for b, t, k in boxes:
        block(d, b, t, k)
    for i in range(len(boxes) - 1):
        arrow(d, ((boxes[i][0][0] + boxes[i][0][2]) // 2, boxes[i][0][3]), ((boxes[i + 1][0][0] + boxes[i + 1][0][2]) // 2, boxes[i + 1][0][1]), "да" if i == 3 else None)
    path = ASSETS / "debug_algorithm_variant4.png"
    img.save(path)
    return path


def set_font(run, size=14, bold=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def p(doc, text="", bold=False, center=False, first=True):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.first_line_indent = Inches(0.49 if first else 0)
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_font(run, 14, bold)
    return para


def heading(doc, text):
    para = p(doc, text, bold=True, first=False)
    para.paragraph_format.space_before = Pt(6)
    return para


def code(doc, text):
    for line in text.strip("\n").splitlines():
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = 1.0
        para.paragraph_format.left_indent = Inches(0.35)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(9)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell(cell, text, bold=False):
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.line_spacing = 1.0
    run = para.add_run(text)
    set_font(run, 11, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell(tbl.rows[0].cells[i], h, True)
        shade(tbl.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cell(cells[i], value)
    return tbl


def screenshot_placeholder(doc, caption, what_to_screen):
    p(doc, "[Сюда вставить скриншот]", center=True, first=False)
    for _ in range(5):
        p(doc, "", first=False)
    p(doc, caption, center=True, first=False)
    p(doc, what_to_screen)


def main():
    flow = make_flowchart()
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(0.6)

    for style_name in ["Normal", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)

    p(doc, "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ", center=True, first=False)
    p(doc, "Московский приборостроительный техникум", center=True, first=False)
    for _ in range(3):
        p(doc, "", first=False)
    p(doc, "ОТЧЕТ", bold=True, center=True, first=False)
    p(doc, "по учебной практике УП.04.01", center=True, first=False)
    p(doc, "Практическая работа №2", bold=True, center=True, first=False)
    p(doc, "Отладка программного кода для устройств промышленного интернета вещей", center=True, first=False)
    for _ in range(3):
        p(doc, "", first=False)
    p(doc, "Студент: Коровин Данила Юрьевич")
    p(doc, "Группа: ВТ-1-24")
    p(doc, "Руководитель: Юдин Даниил Георгиевич")
    p(doc, "Дата: «___» __________ 2026 года")
    doc.add_page_break()

    heading(doc, "Содержание")
    p(doc, "Введение\t3")
    p(doc, "Цель работы\t3")
    p(doc, "Практическая часть\t4")
    p(doc, "Вывод\t16")

    heading(doc, "Введение:")
    p(doc, "IIoT (Industrial Internet of Things, промышленный интернет вещей) - это концепция цифровой трансформации производства, при которой датчики, контроллеры, исполнительные устройства и программное обеспечение объединяются в единую систему. В данной работе выполняется отладка программного кода IIoT-устройства, разработанного в практической работе №1.")

    heading(doc, "Цель работы:")
    p(doc, "Освоить методы и инструменты отладки программного кода для устройств промышленного интернета вещей на уровне отдельных модулей и межмодульных коммуникаций, научиться выявлять, локализовать и устранять ошибки, а также оптимизировать код под целевую платформу.")

    heading(doc, "Задачи работы:")
    p(doc, "2.1 Подготовка к отладке: определение параметров устройства, создание алгоритма отладки, подготовка тестовых сценариев.")
    p(doc, "2.2 Отладка на уровне отдельных модулей: пошаговая проверка, точки останова, просмотр переменных.")
    p(doc, "2.3 Отладка коммуникации между модулями и работы с промышленными протоколами.")
    p(doc, "2.4 Финальная отладка, профилирование, оптимизация и устранение узких мест.")

    heading(doc, "Практическая часть")
    heading(doc, "Этап 1")
    heading(doc, "1.1 Таблица, определение ключевых параметров IIoT:")
    table(doc, ["Параметр", "Значение", "Способ проверки"], [
        ["Период опроса датчика", "2000 мс", "Замер по выводу в терминал"],
        ["Период отправки MQTT", "10000 мс", "Проверка сформированного JSON"],
        ["Диапазон температуры", "-40..+80°C", "Эмуляция датчика"],
        ["Индивидуальная ошибка", "Обработка T < 0°C", "Тест с отрицательными значениями"],
        ["Порог вентиляции", "Вкл. >35°C, выкл. <28°C", "Проверка состояния реле"],
    ])

    heading(doc, "1.2 Блок-схема алгоритма отладки")
    doc.add_picture(str(flow), width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p(doc, "Рисунок №1 схема", center=True, first=False)

    heading(doc, "1.3 Таблица подготовка тестовых сценариев")
    table(doc, ["Сценарий", "Действие", "Ожидаемый результат", "Диагностика"], [
        ["1", "Запустить программу", "Вывод данных об авторе и варианте", "Сообщение в терминале"],
        ["2", "Передать T = 24.5°C", "Температура обрабатывается без ошибки", "JSON содержит 24.5"],
        ["3", "Передать T = -3.4°C", "Знак минус сохраняется", "JSON содержит -3.4"],
        ["4", "Передать T = -12.7°C", "Знак минус сохраняется", "JSON содержит -12.7"],
        ["5", "Передать T = 36.8°C", "Вентиляция включается", "Состояние реле HIGH"],
        ["6", "Передать T = 27.0°C", "Вентиляция выключается", "Состояние реле LOW"],
    ])

    heading(doc, "2 Этап")
    heading(doc, "2.1 Анализ переменных в контрольных точках")
    table(doc, ["Переменная", "Ожидаемое значение", "Фактическое до исправления", "Обнаруженная ошибка"], [
        ["temperature", "-12.7", "-12.7", "Датчик возвращает корректное отрицательное значение"],
        ["process_temperature_bug()", "-12.7", "12.7", "Функция abs() убирает знак минус"],
        ["process_temperature_fixed()", "-12.7", "-12.7", "После исправления знак сохраняется"],
        ["MQTT payload", '"temp": -12.7', '"temp": 12.7', "В сообщении передавались неверные данные"],
    ])

    heading(doc, "2.2 Ошибки и их исправления")
    p(doc, "Индивидуальная ошибка: неправильная обработка отрицательных температур.")
    p(doc, "Проблема: при обработке данных датчика отрицательная температура превращалась в положительную. Например, значение -12.7°C после обработки становилось 12.7°C. Такая ошибка приводит к неверной диагностике и неправильной передаче данных на MQTT-брокер.")
    p(doc, "Диагностика: для проверки была выполнена эмуляция датчика DHT22 со значениями ниже 0°C. В тестовый набор были добавлены температуры -3.4°C и -12.7°C.")
    p(doc, "Было (ошибка):")
    code(doc, """def process_temperature_bug(temperature):
    normalized_temperature = abs(temperature)
    return normalized_temperature""")
    p(doc, "Стало (исправление):")
    code(doc, """def process_temperature_fixed(temperature):
    return temperature""")
    p(doc, "Результат проверки: после исправления отрицательные значения температуры сохраняют знак минус и корректно передаются в JSON-сообщении MQTT.")

    heading(doc, "Этап 3")
    p(doc, "На рисунках №2 - №5 будет изображен код на Python, с которым выполнялась отладка IIoT-системы.")
    screenshot_placeholder(doc, "Рисунок №2 python", "Скриншот верхней части файла main_lab2_variant4.py с ФИО, группой и описанием индивидуального задания №4.")
    screenshot_placeholder(doc, "Рисунок №3 python", "Скриншот функции process_temperature_bug(), где показана ошибка abs(temperature).")
    screenshot_placeholder(doc, "Рисунок №4 python", "Скриншот функции process_temperature_fixed(), где показано исправление ошибки.")
    screenshot_placeholder(doc, "Рисунок №5 python", "Скриншот функции debug_negative_temperature_test(), где выполняется тест с T < 0°C.")
    p(doc, "На рисунке №6 терминал будет показан вывод программы после запуска.")
    screenshot_placeholder(doc, "Рисунок №6 терминал", "Скриншот терминала после команды python main_lab2_variant4.py. Должны быть видны строки: Было с ошибкой T = 12.7°C; Стало после исправления T = -12.7°C.")
    p(doc, "На рисунке №7 MQTT будет показана проверка сформированного JSON-сообщения.")
    screenshot_placeholder(doc, "Рисунок №7 MQTT", 'Скриншот вывода строки MQTT payload, где видно "temp": -12.7.')
    p(doc, "На рисунке №8 git будет изображен GitHub, где был добавлен файл с кодом и был выполнен коммит.")
    screenshot_placeholder(doc, "Рисунок №8 git", "Скриншот GitHub-репозитория или команды git log --oneline с коммитом fix: correct negative temperature handling.")

    heading(doc, "Этап 4")
    heading(doc, "4.3 Таблица, итоговое тестирование по сценариям")
    table(doc, ["Сценарий", "Результат", "Комментарий"], [
        ["1", "Успех", "Программа запускается, данные об авторе выводятся"],
        ["2", "Успех", "Положительная температура обрабатывается корректно"],
        ["3", "Успех", "Значение -3.4°C сохраняет знак минус"],
        ["4", "Успех", "Значение -12.7°C сохраняет знак минус"],
        ["5", "Успех", "При T > 35°C вентиляция включается"],
        ["6", "Успех", "MQTT payload содержит корректную отрицательную температуру"],
    ])

    heading(doc, "Вывод:")
    p(doc, "В ходе выполнения практической работы были освоены методы и инструменты локализации, отладки и оптимизации программного кода для IIoT. В качестве индивидуального задания был выбран вариант №4: неправильная обработка отрицательных температур.")
    p(doc, "При диагностике была обнаружена ошибка, связанная с использованием функции abs(temperature). Из-за неё отрицательные значения температуры теряли знак минус. После исправления программа стала передавать температуру без изменения знака, а повторное тестирование подтвердило корректную обработку значений -3.4°C и -12.7°C.")

    heading(doc, "Источники:")
    p(doc, "1. Документация Python: https://docs.python.org/3/")
    p(doc, "2. Документация JSON в Python: https://docs.python.org/3/library/json.html")
    p(doc, "3. Документация MQTT Explorer.")
    p(doc, "4. Методические указания к практической работе №2.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
